"""
One-time authoring script for samples/ demo fixtures. Not part of the app -
run manually when adding/editing sample content, not on every install.
Needs fpdf2 (pip install fpdf2), not a runtime dependency of the app itself.

Usage: python scripts/generate_samples.py
"""
import os
from pathlib import Path

from docx import Document
from fpdf import FPDF

SAMPLES_DIR = Path(__file__).resolve().parent.parent / "samples"


def write_md(filename: str, content: str) -> None:
    (SAMPLES_DIR / filename).write_text(content.strip() + "\n", encoding="utf-8")


def write_pdf(filename: str, title: str, lines: list[str]) -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, title)
    pdf.set_x(pdf.l_margin)
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 11)
    for line in lines:
        pdf.set_x(pdf.l_margin)
        if line.startswith("## "):
            pdf.ln(2)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 7, line[3:])
            pdf.set_font("Helvetica", "", 11)
        elif line == "":
            pdf.ln(2)
        else:
            pdf.multi_cell(0, 6, line)
    pdf.output(str(SAMPLES_DIR / filename))


def write_docx(filename: str, title: str, sections: list[tuple[str, str]]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for heading, body in sections:
        doc.add_heading(heading, level=2)
        for line in body.strip().split("\n"):
            doc.add_paragraph(line)
    doc.save(str(SAMPLES_DIR / filename))


# ---------------------------------------------------------------------------
# 1. CHF - Jane Doe - DEMO-CHF-01 - 3 files (txt already exists)
# ---------------------------------------------------------------------------

write_md(
    "chf-jane-doe__labs-panel.md",
    """
# Repeat BMP - Jane Doe (DEMO-CHF-01)

MRN 00291847 | Drawn 4h after admission per CHF/hyperkalemia protocol

| Test | Result | Flag |
|---|---|---|
| Na | 138 | normal |
| K | 5.3 | improved from 6.1, still high |
| BUN | 40 | elevated |
| Creatinine | 2.1 | elevated, trending down from 2.3 |
| BNP | pending | - echo ordered, not yet resulted |

**Note:** Lisinopril and spironolactone remain held pending K normalization.
Cardiology re-consulted given persistent hyperkalemia. Echo still pending at
time of this panel - no imaging report available yet.
""",
)

write_docx(
    "chf-jane-doe__discharge-summary.docx",
    "Discharge Summary - Jane Doe (DEMO-CHF-01)",
    [
        ("Admission", "MRN 00291847 | Admitted for acute decompensated CHF exacerbation with hyperkalemia."),
        (
            "Hospital Course",
            "IV Lasix diuresis initiated, net negative 4.2L over 3 days. Potassium normalized to 4.6 "
            "by hospital day 2 after holding lisinopril and spironolactone. Creatinine returned to "
            "baseline 1.8. Echo showed EF 33%, unchanged from prior. Cardiology recommends resuming "
            "lisinopril at reduced dose once outpatient BMP confirms stable potassium.",
        ),
        (
            "Discharge Medications",
            "Furosemide 60mg PO daily (increased)\nCarvedilol 12.5mg PO BID\nSpironolactone HELD - "
            "outpatient cardiology to reassess in 2 weeks\nMetformin 500mg PO BID (resumed)\nAspirin 81mg PO daily",
        ),
        (
            "Follow-up",
            "Cardiology clinic in 1 week. Nephrology in 2 weeks. Repeat BMP at PCP within 5 days. "
            "Daily weights, call clinic if weight gain >3lbs in 2 days.",
        ),
    ],
)

# ---------------------------------------------------------------------------
# 2. Sepsis - John Smith - DEMO-SEP-01 - 2 files (still in ICU, no discharge yet)
# ---------------------------------------------------------------------------

write_pdf(
    "sepsis-john-smith__micro-report.pdf",
    "Microbiology Report - John Smith (DEMO-SEP-01)",
    [
        "MRN 00384921  |  Collected on admission, day 12 hours prior to this report",
        "",
        "## Urine Culture",
        "Result: E. coli, >100,000 CFU/mL",
        "Sensitive to: Ceftriaxone, Piperacillin-Tazobactam, Meropenem",
        "Resistant to: Ampicillin, Trimethoprim-sulfamethoxazole",
        "",
        "## Blood Culture (2 sets)",
        "Set 1: Positive at 14 hours - E. coli, same organism as urine",
        "Set 2: Positive at 16 hours - E. coli, concordant",
        "",
        "## Procalcitonin Trend",
        "Admission: 42 ng/mL",
        "24h recheck: 28 ng/mL (improving on current antibiotics)",
        "",
        "## Interpretation",
        "Gram-negative bacteremia secondary to urosepsis, source confirmed. Current empiric "
        "Piperacillin-Tazobactam + Vancomycin covers isolate. Vancomycin may be discontinued "
        "given gram-negative-only growth pending clinical correlation.",
    ],
)

# ---------------------------------------------------------------------------
# 3. STEMI - Robert Chen - DEMO-MI-01 - 3 files
# ---------------------------------------------------------------------------

write_pdf(
    "stemi-robert-chen__cath-report.pdf",
    "Cardiac Catheterization Report - Robert Chen (DEMO-MI-01)",
    [
        "MRN 00472610  |  Procedure performed emergently on arrival",
        "",
        "## Indication",
        "STEMI, anterior wall, door-to-balloon target <90 min",
        "",
        "## Findings",
        "LAD: 100% occlusion, proximal segment",
        "LCx: 40% stenosis, non-flow-limiting",
        "RCA: mild luminal irregularity, no significant stenosis",
        "",
        "## Intervention",
        "Successful PCI of proximal LAD with drug-eluting stent. TIMI 3 flow restored.",
        "Door-to-balloon time: 62 minutes.",
        "",
        "## Post-Procedure Plan",
        "Dual antiplatelet therapy (Aspirin + Ticagrelor) x 12 months. Continue statin, "
        "beta-blocker, ACE inhibitor once hemodynamically stable. CCU admission overnight.",
    ],
)

write_docx(
    "stemi-robert-chen__discharge-summary.docx",
    "Discharge Summary - Robert Chen (DEMO-MI-01)",
    [
        ("Admission", "MRN 00472610 | Admitted for anterior STEMI, emergent PCI to proximal LAD."),
        (
            "Hospital Course",
            "Uncomplicated post-PCI course. Peak troponin 42 ng/mL. Echo showed EF 45%, mild "
            "anterior wall hypokinesis. No arrhythmias on telemetry. Ambulated without symptoms "
            "by day 2.",
        ),
        (
            "Discharge Medications",
            "Aspirin 81mg PO daily\nTicagrelor 90mg PO BID\nAtorvastatin 80mg PO daily\n"
            "Metoprolol succinate 25mg PO daily\nLisinopril 10mg PO daily (resumed)",
        ),
        (
            "Follow-up",
            "Cardiology in 1 week. Cardiac rehab referral placed. Smoking cessation counseling "
            "reinforced - patient quit 2020, no relapse. Return precautions reviewed.",
        ),
    ],
)

# ---------------------------------------------------------------------------
# 4. DKA - Maria Gonzalez - DEMO-DKA-01 - 3 files
# ---------------------------------------------------------------------------

write_md(
    "dka-maria-gonzalez__icu-flowsheet.md",
    """
# ICU Insulin Drip Flowsheet - Maria Gonzalez (DEMO-DKA-01)

MRN 00519384 | Insulin drip 0.1 units/kg/hr per DKA protocol

| Time | Glucose (mg/dL) | K (mEq/L) | Bicarb | Insulin rate | Notes |
|---|---|---|---|---|---|
| 00:00 (admit) | 512 | 5.9 | 8 | 6 units/hr | Anion gap 28, pH 7.08 |
| 02:00 | 398 | 5.1 | 11 | 6 units/hr | K trending down as expected |
| 04:00 | 310 | 4.3 | 14 | 4 units/hr | K replacement started |
| 06:00 | 245 | 4.0 | 17 | 3 units/hr | Dextrose added to IVF |
| 08:00 | 198 | 4.1 | 19 | 2 units/hr | Anion gap closing |
| 10:00 | 165 | 4.2 | 21 | 1 unit/hr | Tolerating clears, nausea resolved |

**Note:** Transition to subcutaneous insulin once anion gap closed and patient
tolerating oral intake. Endocrinology to guide pump restart.
""",
)

write_docx(
    "dka-maria-gonzalez__discharge-summary.docx",
    "Discharge Summary - Maria Gonzalez (DEMO-DKA-01)",
    [
        ("Admission", "MRN 00519384 | Admitted for severe DKA, precipitated by 48h missed insulin pump doses."),
        (
            "Hospital Course",
            "Insulin drip per DKA protocol, anion gap closed by hour 14. Transitioned to "
            "subcutaneous insulin glargine + lispro. Pump supplies replaced, patient re-educated "
            "on backup insulin regimen for pump failures. No infectious trigger identified.",
        ),
        (
            "Discharge Medications",
            "Insulin glargine 22 units SQ nightly\nInsulin lispro via pump, resumed with new supplies\n"
            "Backup insulin pen prescribed for pump failure scenarios",
        ),
        (
            "Follow-up",
            "Endocrinology in 1 week. Diabetes educator visit scheduled. Discussed sick-day rules "
            "and when to seek care for pump malfunction. Celiac disease diet unchanged.",
        ),
    ],
)

# ---------------------------------------------------------------------------
# 5. Stroke - Dorothy Williams - DEMO-CVA-01 - 4 files (most complex case)
# ---------------------------------------------------------------------------

write_pdf(
    "stroke-dorothy-williams__imaging-report.pdf",
    "Neuroimaging Report - Dorothy Williams (DEMO-CVA-01)",
    [
        "MRN 00603847  |  CT head non-contrast + CT angiography, performed on arrival",
        "",
        "## CT Head (non-contrast)",
        "No acute hemorrhage. Hyperdense right MCA sign, suggestive of acute thrombus.",
        "No established infarct or mass effect at time of scan.",
        "",
        "## CT Angiography Head/Neck",
        "Occlusion of the right M1 segment, middle cerebral artery.",
        "Carotid arteries patent bilaterally, no significant stenosis.",
        "",
        "## Impression",
        "Acute right MCA large vessel occlusion. Findings support tPA administration and "
        "urgent mechanical thrombectomy, consistent with 75-minute last-known-well-to-door time.",
    ],
)

write_md(
    "stroke-dorothy-williams__anticoag-note.md",
    """
# Pharmacy / Neurology Note - Anticoagulation Management

**Patient:** Dorothy Williams (DEMO-CVA-01) | MRN 00603847

## Issue
Patient on warfarin for atrial fibrillation, INR 1.4 on arrival (subtherapeutic).
tPA eligibility requires INR <1.7 - patient qualifies.

## Plan
- Warfarin HELD on admission, do not resume until neurology clears (minimum 24h post-tPA)
- No heparin bridge - contraindicated post-tPA for 24 hours per stroke protocol
- Post-thrombectomy: reassess anticoagulation strategy given large infarct risk of
  hemorrhagic transformation. Consider DOAC vs. warfarin at follow-up given lower
  bleeding risk profile, pending neurology and cardiology input.
- Repeat CT head at 24h to confirm no hemorrhagic transformation before resuming
  any anticoagulation.
""",
)

write_docx(
    "stroke-dorothy-williams__discharge-summary.docx",
    "Discharge Summary - Dorothy Williams (DEMO-CVA-01)",
    [
        ("Admission", "MRN 00603847 | Admitted for acute ischemic stroke, right MCA occlusion, s/p tPA and mechanical thrombectomy."),
        (
            "Hospital Course",
            "IV tPA administered within window, followed by successful mechanical thrombectomy "
            "with TICI 3 reperfusion. NIHSS improved from 14 on arrival to 4 at discharge. "
            "Residual mild left-sided weakness, improving with therapy. No hemorrhagic "
            "transformation on 24h repeat CT.",
        ),
        (
            "Discharge Medications",
            "Apixaban 5mg PO BID (switched from warfarin per cardiology/neurology joint decision)\n"
            "Metoprolol succinate 50mg PO daily\nLisinopril 10mg PO daily (resumed)\n"
            "Atorvastatin 40mg PO daily\nMetformin 500mg PO BID (resumed)",
        ),
        (
            "Follow-up",
            "Neurology in 2 weeks. Inpatient rehab facility for continued PT/OT/speech therapy. "
            "Cardiology to monitor anticoagulation switch. Repeat NIHSS at follow-up visit.",
        ),
    ],
)

print(f"Generated samples in {SAMPLES_DIR}")
for f in sorted(SAMPLES_DIR.glob("*")):
    if f.is_file():
        print(f"  {f.name}")
